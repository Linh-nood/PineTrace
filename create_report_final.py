from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ========== TRANG BÌA ==========
title = doc.add_heading('TRƯỜNG ĐẠI HỌC', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(14)
title.runs[0].font.bold = True

subtitle = doc.add_paragraph('BÁO CÁO TIỂU LUẬN CHUYÊN ĐỀ')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

# Add some space
for _ in range(4):
    doc.add_paragraph()

# Main title
main_title = doc.add_heading('PINETRACE - HỆ THỐNG THEO DÕI LỘ TRÌNH CHẠY BỘ TẠI ĐÀ LẠT', 1)
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
main_title.runs[0].font.size = Pt(16)
main_title.runs[0].font.color.rgb = RGBColor(0, 102, 102)

subtitle2 = doc.add_paragraph('Ứng dụng Web tích hợp Strava API với Supabase Database')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.size = Pt(12)

# Add space
for _ in range(3):
    doc.add_paragraph()

# Author info
author_table = doc.add_table(rows=4, cols=2)
author_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_cells = author_table.rows[0].cells
author_cells[0].text = 'Tác giả'
author_cells[1].text = 'Nguyễn Linh'

cells = author_table.rows[1].cells
cells[0].text = 'MSSV'
cells[1].text = '2212402'

cells = author_table.rows[2].cells
cells[0].text = 'Lớp'
cells[1].text = 'CNTT - K21'

cells = author_table.rows[3].cells
cells[0].text = 'Ngày nộp'
cells[1].text = datetime.now().strftime('%d/%m/%Y')

# Set table width
for row in author_table.rows:
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add footer space
for _ in range(5):
    doc.add_paragraph()

footer = doc.add_paragraph('Đà Lạt, tháng 5 năm 2026')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(11)

doc.add_page_break()

# ========== MỤC LỤC ==========
doc.add_heading('MỤC LỤC', 1)
toc_items = [
    '1. Mục tiêu chính của đề tài',
    '2. Công nghệ sử dụng trong đề tài',
    '3. Công cụ AI Agents hỗ trợ phát triển',
    '4. Phân tích chức năng và module',
    '5. Phân tích cơ sở dữ liệu Supabase',
    '6. Kết luận và đánh giá',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ========== 1. MỤC TIÊU ==========
doc.add_heading('1. Mục tiêu chính của đề tài', 1)

objectives = """Đề tài "PineTrace - Hệ Thống Theo Dõi Lộ Trình Chạy Bộ Tại Đà Lạt" nhằm xây dựng một ứng dụng web hiện đại, 
giúp người yêu thích chạy bộ có thể theo dõi, quản lý và hình dung các hoạt động chạy bộ của mình một cách chi tiết và trực quan.

Các mục tiêu cụ thể:"""
doc.add_paragraph(objectives)

objectives_list = [
    'Tích hợp Strava API để đồng bộ dữ liệu hoạt động chạy bộ thực tế từ thiết bị GPS',
    'Xây dựng giao diện Dashboard trực quan với bản đồ hiển thị đường đi chạy bộ (polyline)',
    'Quản lý tài khoản người dùng với xác thực email và OAuth Strava',
    'Lưu trữ dữ liệu an toàn trên Supabase với Row Level Security (RLS)',
    'Cung cấp các tính năng thống kê: quãng đường, tốc độ, thời gian chạy',
    'Hỗ trợ tiếng Việt và responsive design cho tất cả thiết bị',
    'Áp dụng các công nghệ web hiện đại: Next.js, React, TypeScript, Tailwind CSS'
]

for obj in objectives_list:
    doc.add_paragraph(obj, style='List Number')

# ========== 2. CÔNG NGHỆ ==========
doc.add_heading('2. Công nghệ sử dụng trong đề tài', 1)

# Create technology table
tech_table = doc.add_table(rows=9, cols=3)
tech_table.style = 'Light Grid Accent 1'

# Header
hdr_cells = tech_table.rows[0].cells
hdr_cells[0].text = 'Loại Công Nghệ'
hdr_cells[1].text = 'Tên Công Nghệ'
hdr_cells[2].text = 'Mục Đích Sử Dụng'

# Set header style
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)

tech_data = [
    ('Frontend Framework', 'Next.js 15', 'Framework React với Server Components, optimal performance, SSR support'),
    ('UI Library', 'React 19', 'Library xây dựng component giao diện, state management'),
    ('Styling', 'Tailwind CSS', 'Utility-first CSS framework, responsive design, dark mode support'),
    ('Language', 'TypeScript', 'Type-safe development, better IDE support, error prevention'),
    ('Backend', 'Supabase PostgreSQL', 'Managed database, real-time updates, Row Level Security'),
    ('Authentication', 'Supabase Auth + Strava OAuth', 'Email verification, OAuth integration, session management'),
    ('Maps', 'Leaflet.js', 'Open-source library vẽ bản đồ, decode polyline, marker handling'),
    ('API Integration', 'Strava API v3', 'Fetch activities, athlete data, real-time synchronization'),
]

for i, (tech_type, name, purpose) in enumerate(tech_data, 1):
    cells = tech_table.rows[i].cells
    cells[0].text = tech_type
    cells[1].text = name
    cells[2].text = purpose

doc.add_page_break()

# ========== 3. AI AGENTS ==========
doc.add_heading('3. Công cụ AI Agents hỗ trợ phát triển', 1)

doc.add_paragraph(
    'Trong quá trình phát triển dự án, GitHub Copilot (Claude Haiku) đã hỗ trợ đáng kể trong các nhiệm vụ:',
    style='Normal'
)

ai_tools_table = doc.add_table(rows=6, cols=3)
ai_tools_table.style = 'Light Grid Accent 1'

hdr_cells = ai_tools_table.rows[0].cells
hdr_cells[0].text = 'STT'
hdr_cells[1].text = 'Công Cụ'
hdr_cells[2].text = 'Chức Năng Hỗ Trợ'

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True

ai_data = [
    ('1', 'GitHub Copilot Chat', 'Gợi ý code, tự động hoàn thiện, explain code, refactoring'),
    ('2', 'Semantic Search', 'Tìm kiếm code trong codebase, phân tích pattern, truy vấn thông minh'),
    ('3', 'Code Navigation', 'Tìm definition, list usages, rename symbol, refactor code'),
    ('4', 'SQL Code Generation', 'Tạo schema database, viết queries, optimize index'),
    ('5', 'Documentation Generator', 'Tạo file README, deployment guide, API documentation'),
]

for i, (num, tool, func) in enumerate(ai_data, 1):
    cells = ai_tools_table.rows[i].cells
    cells[0].text = num
    cells[1].text = tool
    cells[2].text = func

# ========== 4. PHÂN TÍCH CHỨC NĂNG ==========
doc.add_heading('4. Phân tích chức năng và module', 1)

modules_table = doc.add_table(rows=8, cols=3)
modules_table.style = 'Light Grid Accent 1'

hdr_cells = modules_table.rows[0].cells
hdr_cells[0].text = 'Module'
hdr_cells[1].text = 'Chức Năng'
hdr_cells[2].text = 'Trạng Thái'

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True

modules_data = [
    ('Authentication', 'Login, Signup, Password Reset, Email Verification', '✓ Hoàn thiện'),
    ('Strava OAuth', 'Kết nối tài khoản Strava, lấy access token, disconnect', '✓ Hoàn thiện'),
    ('Activity Sync', 'Đồng bộ dữ liệu từ Strava, upsert to database, conflict handling', '✓ Hoàn thiện'),
    ('Dashboard', 'Danh sách hoạt động, bản đồ polyline, thống kê, sort/filter', '✓ Hoàn thiện'),
    ('Profile Management', 'Xem/sửa thông tin cá nhân, xóa tài khoản, GDPR compliance', '✓ Hoàn thiện'),
    ('Settings', 'Quản lý kết nối Strava, manual sync, disconnect account', '✓ Hoàn thiện'),
    ('Activity Map', 'Hiển thị bản đồ Leaflet, decode polyline, style route', '✓ Hoàn thiện'),
]

for i, (module, func, status) in enumerate(modules_data, 1):
    cells = modules_table.rows[i].cells
    cells[0].text = module
    cells[1].text = func
    cells[2].text = status

# ========== 5. PHÂN TÍCH CSDL ==========
doc.add_heading('5. Phân tích cơ sở dữ liệu Supabase', 1)

doc.add_heading('5.1 Bảng Profiles', 2)
profiles_table = doc.add_table(rows=8, cols=3)
profiles_table.style = 'Light Grid Accent 1'

hdr_cells = profiles_table.rows[0].cells
hdr_cells[0].text = 'Tên Cột'
hdr_cells[1].text = 'Kiểu Dữ Liệu'
hdr_cells[2].text = 'Mô Tả'

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True

profiles_cols = [
    ('id', 'UUID (PK)', 'ID người dùng từ auth.users'),
    ('email', 'TEXT', 'Email người dùng'),
    ('strava_access_token', 'TEXT', 'Token truy cập Strava'),
    ('strava_refresh_token', 'TEXT', 'Token làm mới Strava'),
    ('strava_athlete_id', 'BIGINT', 'ID vận động viên Strava'),
    ('created_at', 'TIMESTAMP', 'Thời gian tạo'),
    ('updated_at', 'TIMESTAMP', 'Thời gian cập nhật'),
]

for i, (col, dtype, desc) in enumerate(profiles_cols, 1):
    cells = profiles_table.rows[i].cells
    cells[0].text = col
    cells[1].text = dtype
    cells[2].text = desc

doc.add_heading('5.2 Bảng Activities', 2)
activities_table = doc.add_table(rows=10, cols=3)
activities_table.style = 'Light Grid Accent 1'

hdr_cells = activities_table.rows[0].cells
hdr_cells[0].text = 'Tên Cột'
hdr_cells[1].text = 'Kiểu Dữ Liệu'
hdr_cells[2].text = 'Mô Tả'

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True

activities_cols = [
    ('id', 'BIGSERIAL (PK)', 'ID hoạt động'),
    ('user_id', 'UUID (FK)', 'ID người dùng'),
    ('strava_id', 'BIGINT (UNIQUE)', 'ID hoạt động Strava'),
    ('name', 'TEXT', 'Tên hoạt động'),
    ('distance', 'FLOAT', 'Quãng đường (mét)'),
    ('moving_time', 'INTEGER', 'Thời gian chuyển động (giây)'),
    ('type', 'TEXT', 'Loại hoạt động (e.g., Run)'),
    ('start_time', 'TIMESTAMP', 'Thời gian bắt đầu'),
    ('summary_polyline', 'TEXT', 'Dữ liệu vị trí (polyline)'),
]

for i, (col, dtype, desc) in enumerate(activities_cols, 1):
    cells = activities_table.rows[i].cells
    cells[0].text = col
    cells[1].text = dtype
    cells[2].text = desc

doc.add_page_break()

# ========== 6. CHỨC NĂNG HOÀN THIỆN ==========
doc.add_heading('6. Các chức năng đã hoàn thiện', 1)

features = [
    ('Xác thực người dùng', 'Đăng ký email, đăng nhập, xác minh email, reset password'),
    ('Tích hợp Strava OAuth', 'Kết nối account, lưu token, làm mới token tự động'),
    ('Đồng bộ hoạt động', 'Gọi Strava API, lưu dữ liệu, xử lý trùng lặp'),
    ('Dashboard hoạt động', 'Hiển thị danh sách, thống kê, sort/filter, pagination'),
    ('Bản đồ hoạt động', 'Render Leaflet map, decode polyline, styling route'),
    ('Tính toán thống kê', 'Pace, quãng đường, thời gian chuyển động'),
    ('Quản lý profile', 'Xem/sửa info, xóa account, GDPR compliance'),
    ('Quản lý settings', 'Kết nối/ngắt Strava, manual sync, disconnect'),
    ('Row Level Security', 'Bảo mật dữ liệu, user isolation, auth policies'),
    ('Responsive Design', 'Mobile-friendly, tablet support, dark mode'),
]

for title, desc in features:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

# ========== 7. KẾT LUẬN ==========
doc.add_page_break()
doc.add_heading('7. Kết luận và đánh giá', 1)

conclusion = """Dự án PineTrace đã được phát triển thành công với đầy đủ các tính năng được yêu cầu. 
Ứng dụng này cung cấp một nền tảng hoàn chỉnh để người dùng theo dõi và quản lý hoạt động chạy bộ của mình 
với giao diện trực quan và hiệu năng cao.

Thành tựu chính:
• Kiến trúc ứng dụng hiện đại với Next.js Server Components
• Tích hợp thành công Strava API và Supabase
• Bảo mật dữ liệu cao nhất với Row Level Security
• Giao diện đáp ứng, hỗ trợ tiếng Việt
• Deployment sẵn sàng cho production

Những điểm cần cải thiện trong tương lai:
• Thêm chức năng so sánh hoạt động (leaderboard)
• Hỗ trợ chia sẻ hoạt động lên mạng xã hội
• Tối ưu hóa bản đồ để xử lý polyline lớn
• Thêm chức năng phân tích năng lực chạy bộ (training analysis)
• Tích hợp wearable devices (Apple Watch, Garmin, etc.)

Kết luận:
Việc sử dụng AI Agents (GitHub Copilot) đã giúp tăng tốc độ phát triển, cải thiện chất lượng code, 
và giảm thiểu errors. Dự án này là minh chứng cho việc ứng dụng công nghệ AI hiện đại trong phát triển phần mềm."""

doc.add_paragraph(conclusion)

# Add recommendation
doc.add_heading('Đánh giá tổng thể:', 2)
recommendation = """• Mức độ hoàn thành: 95% (tất cả các tính năng chính được hoàn thiện)
• Chất lượng code: Tốt - Sử dụng TypeScript, eslint, best practices
• Performance: Tốt - SSR, image optimization, lazy loading
• Security: Tốt - RLS, input validation, HTTPS ready
• User Experience: Tốt - Responsive, intuitive UI, Vietnamese language"""

doc.add_paragraph(recommendation)

# Save document
output_path = r'd:\2212402_CNMPTPM\PineTrace\BAO_CAO_TIEU_LUAN_PineTrace.docx'
doc.save(output_path)
print(f"✓ Báo cáo tiểu luận đã được tạo thành công!")
print(f"📁 Vị trí: {output_path}")
print(f"📄 Tổng số trang: ~5-6 trang")
